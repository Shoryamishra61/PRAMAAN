"""Populate the retained System Design DOCX without replacing its visual system."""

from __future__ import annotations

import hashlib
import os
import shutil
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
REFERENCE = Path(
    r"C:\Users\SHORYA MISHRA\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.1\skills\artifact-template-system-design\assets\reference.docx"
)
OUTPUT = ROOT / "output" / "CARVE-System-Design.docx"
REFERENCE_SHA256 = "13504f6c221a42c1726460a9e865e563355539ff97d702d6c9b2267b4b261d76"


def replace_paragraph(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph(paragraph, text)
    for extra in cell.paragraphs[1:]:
        replace_paragraph(extra, "")


def fill_table(table, rows: list[list[str]]) -> None:
    if len(rows) != len(table.rows):
        raise ValueError("Template table row count changed.")
    for row, values in zip(table.rows, rows, strict=True):
        if len(values) != len(row.cells):
            raise ValueError("Template table column count changed.")
        for cell, value in zip(row.cells, values, strict=True):
            set_cell(cell, value)


def architecture_diagram() -> bytes:
    width, height = 1568, 800
    image = Image.new("RGB", (width, height), "#f7f9fb")
    draw = ImageDraw.Draw(image)
    regular_path = Path(r"C:\Windows\Fonts\segoeui.ttf")
    bold_path = Path(r"C:\Windows\Fonts\segoeuib.ttf")
    regular = ImageFont.truetype(str(regular_path), 23)
    small = ImageFont.truetype(str(regular_path), 18)
    bold = ImageFont.truetype(str(bold_path), 25)
    title = ImageFont.truetype(str(bold_path), 30)
    navy, blue, pale, green, amber, red = (
        "#0a3155",
        "#2f73a8",
        "#e7f1f8",
        "#19704d",
        "#a96800",
        "#a63131",
    )
    draw.rectangle((0, 0, width, 108), fill=navy)
    draw.text((46, 26), "CARVE AUTHORITY BOUNDARY", font=title, fill="white")
    draw.text(
        (46, 70),
        "Evidence to grounded relation to formal proof to safe decision",
        font=small,
        fill="#c7d9e8",
    )

    stages = [
        ("Evidence ingestion", "Provenance + schema", blue),
        ("Semantic relations", "Exact grounded spans", blue),
        ("Financial proof", "Immutable invariants", green),
        ("Selective control", "Risk / OOD abstention", amber),
        ("PASS / REVIEW / BLOCK", "Certificate + diff", navy),
    ]
    x0, y0, box_w, box_h, gap = 45, 185, 250, 148, 53
    for index, (label, subtitle, color) in enumerate(stages):
        left = x0 + index * (box_w + gap)
        draw.rounded_rectangle(
            (left, y0, left + box_w, y0 + box_h), radius=12, fill="white", outline=color, width=4
        )
        draw.rectangle((left, y0, left + box_w, y0 + 15), fill=color)
        draw.text((left + 18, y0 + 38), label, font=bold, fill=navy)
        draw.text((left + 18, y0 + 88), subtitle, font=small, fill="#42576a")
        if index < len(stages) - 1:
            start = (left + box_w + 8, y0 + box_h // 2)
            end = (left + box_w + gap - 8, y0 + box_h // 2)
            draw.line((start, end), fill=navy, width=5)
            draw.polygon(
                ((end[0], end[1]), (end[0] - 15, end[1] - 10), (end[0] - 15, end[1] + 10)),
                fill=navy,
            )

    lower = [
        (
            92,
            448,
            565,
            612,
            "Active evidence acquisition",
            "Request minimum-cost missing authoritative evidence",
            pale,
            amber,
        ),
        (
            605,
            448,
            1098,
            612,
            "Append-only audit + replay",
            "Digests, versions, proof trace, decision, repair",
            pale,
            blue,
        ),
        (
            1138,
            448,
            1515,
            612,
            "Human authority",
            "Inspect, repair, or locally override",
            "#e7f3ec",
            green,
        ),
    ]
    for left, top, right, bottom, label, subtitle, fill, color in lower:
        draw.rounded_rectangle(
            (left, top, right, bottom), radius=12, fill=fill, outline=color, width=3
        )
        draw.text((left + 20, top + 30), label, font=bold, fill=navy)
        draw.multiline_text((left + 20, top + 78), subtitle, font=small, fill="#42576a", spacing=7)

    draw.line((217, 448, 217, 382), fill=amber, width=4)
    draw.polygon(((217, 370), (206, 388), (228, 388)), fill=amber)
    draw.line((850, 333, 850, 448), fill=blue, width=4)
    draw.polygon(((850, 448), (839, 430), (861, 430)), fill=blue)
    draw.line((1326, 333, 1326, 448), fill=green, width=4)
    draw.polygon(((1326, 448), (1315, 430), (1337, 430)), fill=green)

    draw.rectangle((45, 682, 1523, 748), fill="#f9e2df", outline=red, width=2)
    draw.text((70, 701), "HARD SAFETY INVARIANT", font=bold, fill=red)
    draw.text(
        (455, 705),
        "No learned component overrides authoritative state. No payment or dispute write exists.",
        font=regular,
        fill="#522b2d",
    )
    buffer = BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    return buffer.getvalue()


def patch_package(path: Path) -> None:
    temp = path.with_suffix(".tmp.docx")
    diagram = architecture_diagram()
    with ZipFile(path, "r") as source, ZipFile(temp, "w", ZIP_DEFLATED) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "word/media/image1.png":
                data = diagram
            elif info.filename == "word/footer1.xml":
                data = data.replace(
                    b"[Organization Name] | System Design RFC",
                    b"Dispute Integrity Gate | CARVE System Design",
                )
            target.writestr(info, data)
    os.replace(temp, path)


def main() -> None:
    digest = hashlib.sha256(REFERENCE.read_bytes()).hexdigest()
    if digest != REFERENCE_SHA256:
        raise RuntimeError("Retained System Design template changed; redistill before authoring.")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(REFERENCE, OUTPUT)
    document = Document(OUTPUT)
    paragraphs = document.paragraphs

    replacements = {
        8: "Dispute Integrity Gate / FECL",
        9: "CARVE: Cost-aware, Risk-controlled Financial Evidence Verification",
        22: (
            "CARVE is a defense-only verifier for refund/credit-not-processed disputes. It "
            "turns untrusted communication into exact grounded claims, reconciles those claims "
            "against authoritative payment/refund state, compiles supported financial invariants, "
            "and returns PASS, REVIEW, or a provenance-grounded BLOCK certificate."
        ),
        23: (
            "The synthetic research system supports interactive evidence import, deliberate "
            "failure cases, active acquisition of missing authoritative evidence, replayable "
            "artifacts, and human repair. It does not submit disputes, move money, predict dispute "
            "wins, claim production prevalence, or let learned models override deterministic truth."
        ),
        28: (
            "Refund disputes can contain plausible but mutually inconsistent emails, refund exports, "
            "amounts, references, currencies, and timestamps. Generic summaries obscure causality; "
            "manual checking is slow; and missing evidence can be mistaken for absence. The narrow "
            "loss is merchant exposure caused by submitting internally inconsistent evidence."
        ),
        29: (
            "The boundary is a read-only merchant-side pre-submission gate. Evidence ingestion and "
            "semantic extraction are untrusted. Immutable structured state and compiled invariants "
            "own financial truth. Unsupported, OOD, malformed, incomplete, or unavailable inputs fail "
            "closed to REVIEW."
        ),
        33: "Figure 1. CARVE authority boundary and evidence-to-decision flow.",
        40: "1. A synthetic case bundle or local JSON/TXT input enters the evidence boundary with provenance metadata.",
        41: "2. Schema, money precision, source type, digest, scope, and exact-span grounding are validated.",
        42: "3. The reason profile, authoritative payment/refund snapshot, frozen model artifacts, and risk policy are loaded by version and digest.",
        43: "4. The semantic layer nominates typed relations; deterministic reconciliation and the proof compiler decide whether supported constraints are SAT, UNSAT, or INCOMPLETE.",
        44: "5. The run digest, grounded spans, constraint trace, decision, acquisition request, and certificate are appended to local audit output before analyst handoff.",
        45: "6. If evidence is incomplete, CARVE requests the minimum-cost supported evidence item. Timeouts, outages, OOD inputs, or invalid artifacts terminate in REVIEW.",
        46: "7. The UI renders the exact evidence, decision authority, before/after repair diff, and zero-network-write boundary; generated metrics remain isolated under /evaluation and /research.",
        54: "Every decision-relevant semantic claim must be schema-valid and resolve to an exact source span before deterministic verification.",
        55: "Runs use content digests and stable case/evidence identifiers; local replays are idempotent and perform no external financial action.",
        56: "Dataset, split, model, threshold, calibration, code, request, evidence, proof, and result digests are captured for audit and replay.",
        57: "Learned scores are advisory residual-risk evidence only; authoritative payment/refund state and failed financial invariants remain the source of truth.",
        58: "Versioned contracts are maintained in backend/app/carve.py, backend/app/sandbox_api.py, and data/financial-evidence-integrity/v4.5/.",
        59: "Any schema or supported-relation change requires an explicit benchmark version bump and a new freeze protocol.",
        63: (
            "Evidence digests protect identity relative to recorded content, not source authenticity. "
            "Family-grouped splits keep causal pairs together. The frozen TEST receipt prevents "
            "reruns after threshold selection. Duplicate local runs produce stable request digests; "
            "partial proof state is REVIEW, never an inferred PASS."
        ),
        67: "Webhook ingestion uses raw-body HMAC and event identifiers; the live research instrument accepts only bounded local synthetic inputs.",
        68: "Evidence text is treated as untrusted and remains local in the demo; logs must exclude secrets and avoid claiming full de-identification.",
        69: "Models receive no payment credentials, database credentials, write tools, or unrestricted network access.",
        70: "Replay, debugging, and acquisition default to read-only behavior; malformed artifacts, digest failures, and provider outages route to REVIEW.",
        71: "Production retention, deletion, residency, tenant isolation, and consent policy remain deployment prerequisites and are not claimed by the synthetic prototype.",
        81: "Which governed real-dispute corpus can support external validation?",
        82: "Which authoritative read surfaces can production acquisition use?",
        83: "Which prevalence and rupee cost matrix should set risk thresholds?",
        84: "Who owns model risk, incidents, drift, and override policy?",
        87: (
            "Retain deterministic proof as the decision authority and learned relations only in their "
            "measured narrow role. Next: governed real-data validation, browser/accessibility QA, and "
            "a safety-approved read-only shadow deployment."
        ),
    }
    for index, text in replacements.items():
        replace_paragraph(paragraphs[index], text)
    # This template slot contains a hyperlink run that python-docx does not expose in
    # paragraph.runs. Reassigning the paragraph text removes the retained placeholder.
    paragraphs[59].text = replacements[59]

    tables = document.tables
    fill_table(
        tables[0],
        [
            [
                "STATUS\nImplemented synthetic research prototype",
                "",
                "OWNER\nDispute Integrity Gate team",
                "",
                "LAST UPDATED\nSeptember 2, 2026",
            ]
        ],
    )
    fill_table(
        tables[1],
        [
            ["Authors", "Dispute Integrity Gate team"],
            ["Reviewers", "Razorpay AI Buildathon judges; fintech risk and ML reviewers"],
            [
                "Related docs",
                "docs/00-SOURCE-OF-TRUTH.md; docs/05-PRD.md; docs/06-SRS.md; CARVE research paper",
            ],
            ["Scope", "Read-only refund-not-processed financial evidence consistency verification"],
        ],
    )
    fill_table(
        tables[2],
        [
            ["Goals", "Non-goals"],
            [
                "Ground every learned claim to an exact source span",
                "Predict dispute wins or legal validity",
            ],
            [
                "Keep authoritative financial invariants outside models",
                "Automate accept, contest, refund, or payment actions",
            ],
            [
                "Measure false-PASS exposure and safe abstention",
                "Claim production performance from synthetic data",
            ],
            [
                "Make evidence acquisition and repair auditable",
                "Support every dispute reason, language, document, or network",
            ],
        ],
    )
    fill_table(
        tables[3],
        [
            ["Component", "Responsibility", "Primary storage", "Failure behavior"],
            [
                "Evidence ingestion/provenance",
                "Validate bounded input, identifiers, digests, and source metadata",
                "Ephemeral request plus append-only local audit",
                "Malformed or unsupported input returns REVIEW/rejection",
            ],
            [
                "Semantic relation extractor",
                "Nominate typed claims with exact source quotes and spans",
                "Frozen model artifact plus versioned schema",
                "Outage, ungrounded output, or OOD returns REVIEW",
            ],
            [
                "Deterministic verifier/proof compiler",
                "Reconcile amounts, references, currency, state, and time; emit certificate",
                "Authoritative snapshot plus constraint trace",
                "Incomplete state returns REVIEW; proven conflict returns BLOCK",
            ],
            [
                "Residual risk/selective controller",
                "Compare frozen research candidates and abstain under unsupported risk",
                "Model/calibration artifacts",
                "Cannot override hard truth; invalid certificate returns REVIEW",
            ],
            [
                "Evidence acquisition/audit UI",
                "Request minimum supported missing evidence; show repair and decision diff",
                "Local sample bundles and artifact hashes",
                "No network or money write; acquisition failure remains REVIEW",
            ],
        ],
    )
    fill_table(
        tables[4],
        [
            ["Field", "Type", "Required", "Description"],
            [
                "raw_reason_code",
                "string",
                "Yes",
                "Preserved verbatim; bounded to the supported local reason profile",
            ],
            [
                "customer_communication",
                "string",
                "Yes",
                "Untrusted semantic source; exact quotes and offsets required",
            ],
            [
                "payment_amount_inr",
                "decimal string",
                "Yes",
                "Parsed deterministically into integer minor units",
            ],
            [
                "refund_ledger_complete",
                "boolean",
                "Yes",
                "Controls whether absence can be treated as authoritative",
            ],
            [
                "refund_status",
                "enum",
                "Yes",
                "Local refund lifecycle state; never inferred by a model",
            ],
            [
                "refund_amount_inr",
                "decimal string/null",
                "Conditional",
                "Required when a refund record exists; at most 2 decimals",
            ],
            [
                "simulation",
                "enum",
                "No",
                "Bounded failure injection for local model-outage testing",
            ],
        ],
    )
    fill_table(
        tables[5],
        [
            ["Scenario", "Expected behavior", "Reasoning"],
            [
                "Duplicate local request",
                "Stable request digest; repeatable result; no side effect",
                "The sandbox is ephemeral and read-only",
            ],
            [
                "Missing authoritative refund export",
                "REVIEW plus REQUEST_REFUND_EXPORT cost 1",
                "Missing state cannot prove agreement or contradiction",
            ],
            [
                "Model outage, OOD, or malformed semantics",
                "REVIEW or input rejection; never PASS",
                "Degraded semantics cannot silently clear the gate",
            ],
            [
                "Artifact or policy changes during execution",
                "Use the frozen version/digest captured at run start",
                "Mid-run authority changes would break replayability",
            ],
        ],
    )
    fill_table(
        tables[6],
        [
            ["Signal", "SLO or alert", "Owner", "Launch gate"],
            [
                "False-PASS count and rupee exposure",
                "0 on frozen synthetic TEST; production target NOT YET MEASURED",
                "Risk/ML",
                "Required",
            ],
            [
                "End-to-end latency",
                "Artifact-backed p50/p95 reported; production SLO NOT YET MEASURED",
                "ML platform",
                "Required before production",
            ],
            [
                "REVIEW, outage, and OOD rate",
                "Slice and trend monitoring; production thresholds NOT YET MEASURED",
                "Risk operations",
                "Required before shadow launch",
            ],
            [
                "Dataset/model/policy drift",
                "Digest mismatch fails closed; drift thresholds need governed real data",
                "Model risk",
                "Required",
            ],
            [
                "Replay and certificate integrity",
                "Every frozen release reruns integrity tests; TEST receipt remains one-shot",
                "Engineering",
                "Required",
            ],
            [
                "Rollout constraint: synthetic demo only. Real-data validation, security review, provenance integration, browser/accessibility QA, rollback rehearsal, and owner approval are required before read-only shadow deployment.",
                "Rollout constraint: synthetic demo only. Real-data validation, security review, provenance integration, browser/accessibility QA, rollback rehearsal, and owner approval are required before read-only shadow deployment.",
                "Rollout constraint: synthetic demo only. Real-data validation, security review, provenance integration, browser/accessibility QA, rollback rehearsal, and owner approval are required before read-only shadow deployment.",
                "Rollout constraint: synthetic demo only. Real-data validation, security review, provenance integration, browser/accessibility QA, rollback rehearsal, and owner approval are required before read-only shadow deployment.",
            ],
        ],
    )
    fill_table(
        tables[7],
        [
            ["Alternative", "Why it was considered", "Why it was not selected"],
            [
                "Rules-only verifier",
                "Perfect frozen financial reconciliation and simplest authority boundary",
                "Retained as selected decision baseline, but cannot generalize semantic relation extraction alone",
            ],
            [
                "End-to-end LLM decision",
                "Flexible language handling",
                "Ungrounded outputs, calibration limits, cost, and unacceptable authority leakage",
            ],
            [
                "Model-stacked ensemble",
                "Potential aggregate predictive lift",
                "No justified lift over simpler deterministic truth and increased failure surface",
            ],
            [
                "Autonomous dispute agent",
                "Could reduce operator steps",
                "Outside defense-only scope and unsafe for consequential financial actions",
            ],
        ],
    )
    fill_table(
        tables[8],
        [
            ["Milestone", "Deliverable", "Exit criteria"],
            [
                "M1",
                "Frozen synthetic CARVE and live debugger",
                "Integrity, one-shot receipt, sample replay, and no-write gates pass",
            ],
            [
                "M2",
                "Governed real-data validation",
                "Blind holdout, annotation agreement, intervals, and cost review pass",
            ],
            [
                "M3",
                "Read-only shadow",
                "Provenance, monitoring, rollback, accessibility, and incident gates pass",
            ],
            [
                "M4",
                "Limited analyst pilot",
                "Stable safety metrics, human-factors evidence, and model-risk approval",
            ],
        ],
    )

    document.core_properties.title = "CARVE System Design"
    document.core_properties.subject = "Defense-only financial evidence consistency verification"
    document.core_properties.author = "Dispute Integrity Gate team"
    document.save(OUTPUT)
    patch_package(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
